"""KALESS Engine — Docx Generator.

Produces professional, high-fidelity APA 7 formatted academic reports using python-docx.
"""
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.export.pdf_generator import _generate_matplotlib_chart

def set_cell_border(cell, **kwargs):
    """Set cell borders in python-docx table for strict APA horizontal line compliance."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def clear_cell_borders(cell):
    """Remove all borders from a cell (APA tables only have horizontal lines)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement('w:{}'.format(edge))
        element.set(qn('w:val'), 'nil')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def generate_docx(result: dict) -> bytes:
    """Generate a premium APA 7 formatted Docx report from statistical results."""
    doc = Document()
    
    # 1. Global Document Styling (APA 7: Times New Roman 12pt, Double Spaced)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0  # APA standard double spacing
    style.paragraph_format.space_after = Pt(0)
    
    # 2. Cover Section (Centered Title)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(result.get('title', 'Statistical Analysis Report').upper())
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    doc.add_paragraph() # Spacer
    
    # Analysis Summary
    summary_p = doc.add_paragraph()
    summary_p.add_run("Kaless Analysis Engine Output").bold = True
    
    timestamp = result.get("metadata", {}).get("timestamp", datetime.now().isoformat())
    try:
        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
        date_str = dt.strftime("%B %d, %Y")
    except:
        date_str = timestamp
        
    doc.add_paragraph(f"Date: {date_str}")
    
    # 3. Iterate Output Blocks
    blocks = result.get("output_blocks", [])
    table_counter = 1
    
    for block in blocks:
        b_type = block.get("block_type")
        b_title = block.get("title", "Statistical Output")
        content = block.get("content", {})
        
        if b_type == "table":
            columns = content.get("columns", [])
            rows = content.get("rows", [])
            
            if not columns or not rows:
                continue
            
            # --- APA 7 Table Label ---
            label_p = doc.add_paragraph()
            label_p.paragraph_format.line_spacing = 1.0
            label_run = label_p.add_run(f"Table {table_counter}")
            label_run.bold = True
            
            # --- APA 7 Table Title ---
            title_p = doc.add_paragraph()
            title_p.paragraph_format.line_spacing = 1.0
            title_run = title_p.add_run(b_title)
            title_run.italic = True
            
            table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
            table.style = 'Table Grid' # Base style
            table.autofit = True
            
            # Remove all default borders first
            for r in table.rows:
                for c in r.cells:
                    clear_cell_borders(c)
            
            # Header Formatting
            hdr_cells = table.rows[0].cells
            for idx, col in enumerate(columns):
                hdr_cells[idx].text = str(col)
                # Bold header? Optional in APA but common in software reports
                p = hdr_cells[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
                p.paragraph_format.line_spacing = 1.0
                
                # Header Horizontal Lines (Top and Bottom)
                set_cell_border(hdr_cells[idx], 
                    top={"sz": 4, "val": "single", "color": "000000"}, 
                    bottom={"sz": 4, "val": "single", "color": "000000"}
                )
                
            # Rows Formatting
            for r_idx, row_data in enumerate(rows):
                row_cells = table.rows[r_idx + 1].cells
                for c_idx, col in enumerate(columns):
                    val = row_data.get(col, "")
                    row_cells[c_idx].text = str(val) if val is not None else "."
                    p = row_cells[c_idx].paragraphs[0]
                    p.paragraph_format.line_spacing = 1.0
                    
                    # Bottom Horizontal Line for last row
                    if r_idx == len(rows) - 1:
                        set_cell_border(row_cells[c_idx], 
                            bottom={"sz": 4, "val": "single", "color": "000000"}
                        )
            
            # --- APA 7 Table Notes ---
            footnotes = content.get("footnotes", [])
            if footnotes:
                note_p = doc.add_paragraph()
                note_p.paragraph_format.line_spacing = 1.0
                note_run = note_p.add_run("Note. ")
                note_run.italic = True
                note_p.add_run("; ".join(footnotes))
            
            doc.add_paragraph() # Spacer after table
            table_counter += 1
                
        elif b_type == "chart":
            doc.add_heading(b_title, level=2)
            chart_type = content.get("chart_type", "bar")
            data = content.get("data", [])
            config = content.get("config", {})
            
            try:
                img_buf = _generate_matplotlib_chart(chart_type, data, config)
                doc.add_picture(img_buf, width=Inches(5.5))
            except Exception as e:
                doc.add_paragraph(f"[Chart generation failed: {str(e)}]")
                
        elif b_type == "text":
            doc.add_paragraph(content.get("text", ""))
            
    # 4. Academic Interpretation (APA write-up)
    interp = result.get("interpretation")
    if interp and interp.get("academic_sentence"):
        doc.add_page_break()
        h = doc.add_heading("Result Interpretation", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Indented and double spaced as per APA
        p = doc.add_paragraph(interp.get("academic_sentence", ""))
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 2.0
        
    # 5. Professional Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run(f"Generated by Kaless Advanced Statistics Engine\nReport Hash: {os.urandom(4).hex().upper()}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
